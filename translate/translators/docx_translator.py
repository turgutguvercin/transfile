from docx import Document
from .utils import batched, client, MODEL, PARA_DELIM, RUN_DELIM

def get_run_texts(paragraph):
    """Return the list of run texts for a single paragraph."""
    return [r.text for r in paragraph.runs]

def set_run_texts(paragraph, new_texts):
    """Write translated texts back into runs without touching styling."""
    runs = paragraph.runs
    k = min(len(runs), len(new_texts))
    for i in range(k):
        runs[i].text = new_texts[i]
    # clear extra original runs if model returned fewer
    for i in range(k, len(runs)):
        runs[i].text = ""
    # if model returned more segments than original runs, append extras to the last run
    if len(new_texts) > k and k > 0:
        runs[k - 1].text += "".join(new_texts[k:])

def translate_docx_chunk(paragraph_payloads, src, tgt):
    """
    paragraph_payloads: list[str], each is one paragraph's runs joined by RUN_DELIM
    Returns list[str] translated, still joined by RUN_DELIM (one string per paragraph)
    """
    if not paragraph_payloads:
        return []

    prompt = (
        f"You are a professional translator. Translate from {src} to {tgt}. "
        f"Crucially, keep ALL delimiters EXACTLY: paragraph delimiter {PARA_DELIM} "
        f"and run delimiter {RUN_DELIM}. Do NOT add or remove delimiters; "
        f"preserve their count. Keep numbers and punctuation intact."
    )
    joined = PARA_DELIM.join(paragraph_payloads)

    r = client.chat.completions.create(
        model=MODEL,
        temperature=0.1,
        messages=[{
            "role": "user",
            "content": f"{prompt}\n\n{joined}"
        }],
    )
    out = (r.choices[0].message.content or "").strip()
    parts = out.split(PARA_DELIM)
    if len(parts) != len(paragraph_payloads):
        raise RuntimeError(f"[DOCX] Paragraph mismatch: sent {len(paragraph_payloads)} got {len(parts)}")
    return parts

def translate_docx(in_path: str, out_path: str, src: str, tgt: str, batch_size: int = 40):
    """Translate a DOCX file preserving formatting"""
    doc = Document(str(in_path))

    # Collect all paragraphs, including those inside table cells, uniformly
    paragraphs = []
    # body paragraphs
    paragraphs.extend([p for p in doc.paragraphs if p.text.strip()])
    # table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p)

    if not paragraphs:
        doc.save(str(out_path))
        print(f"No translatable paragraphs. Saved unchanged → {out_path}")
        return

    # Prepare payloads: each paragraph as runs joined with RUN_DELIM
    items = []
    for p in paragraphs:
        runs = get_run_texts(p)
        # normalize empty runs so run count stays stable
        norm = [t if t != "" else " " for t in runs]
        items.append(RUN_DELIM.join(norm))

    # Translate in batches (keep order)
    out_items = []
    for chunk in batched(items, batch_size):
        out_items.extend(translate_docx_chunk(chunk, src, tgt))

    # Write back run-by-run
    for p, para_txt in zip(paragraphs, out_items):
        run_texts = para_txt.split(RUN_DELIM)
        set_run_texts(p, run_texts)

    doc.save(str(out_path))
    print(f"Translated DOCX → {out_path}")
